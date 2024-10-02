import io
import os
import re
import json
import time
import boto3
import base64
import PyPDF2
import string
import random
from docx import Document
import pdfplumber
import pandas as pd
import streamlit as st
from docx import Document
from requests import request
from sigv4 import SigV4HttpRequester

# AWS Session and Clients Instantiation
region = os.environ['AWS_REGION']
session = boto3.Session(region_name=region)
agent_client = session.client('bedrock-agent')
agent_runtime_client = session.client('bedrock-agent-runtime')
s3_client = session.client('s3')

# Streamlit Agents and Knowledge Bases Helper Functions
def generate_session_id():
    return ''.join(random.choices(string.digits, k=5))

# Function to fetch agent data
def fetch_agents():
    agents = []
    next_token = None
    while True:
        response = agent_client.list_agents(maxResults=10, nextToken=next_token) if next_token else agent_client.list_agents(maxResults=10)
        agents.extend(response['agentSummaries'])
        next_token = response.get('nextToken')
        if not next_token:
            break
    return agents

# Function to fetch knowledge base data
def fetch_knowledge_bases():
    knowledge_bases = []
    next_token = None
    while True:
        response = agent_client.list_knowledge_bases(maxResults=10, nextToken=next_token) if next_token else agent_client.list_knowledge_bases(maxResults=10)
        knowledge_bases.extend(response['knowledgeBaseSummaries'])
        next_token = response.get('nextToken')
        if not next_token:
            break
    return knowledge_bases

# Function to fetch data sources and their IDs and names
def fetch_data_sources(kb_id):
    data_sources = []
    next_token = None
    while True:
        response = agent_client.list_data_sources(knowledgeBaseId=kb_id, maxResults=100, nextToken=next_token) if next_token else agent_client.list_data_sources(knowledgeBaseId=kb_id, maxResults=100)
        for ds in response['dataSourceSummaries']:
            data_source_info = {
                'id': ds['dataSourceId'],
                'name': ds['name']
            }
            data_sources.append(data_source_info)
        next_token = response.get('nextToken')
        if not next_token:
            break
    return data_sources

# Function to fetch agent aliases
def fetch_agent_aliases(agent_id):
    agent_aliases = []
    next_token = None
    while True:
        response = agent_client.list_agent_aliases(agentId=agent_id, maxResults=5, nextToken=next_token) if next_token else agent_client.list_agent_aliases(agentId=agent_id, maxResults=5)
        agent_aliases.extend(response['agentAliasSummaries'])
        next_token = response.get('nextToken')
        if not next_token:
            break
    return agent_aliases

# Function to list knowledge bases associated with an agent
def fetch_agent_knowledge_bases(agent_id):
    response = agent_client.list_agent_knowledge_bases(
        agentId=agent_id,
        agentVersion='DRAFT'
    )
    return response['agentKnowledgeBaseSummaries']

# Function to get knowledge base name
def fetch_knowledge_base_name(kb_id):
    response = agent_client.get_knowledge_base(knowledgeBaseId=kb_id)
    return response['knowledgeBase']['name']

# Function to S3 bucket name
def extract_bucket_name(bucket_arn):
    # ARN format: arn:aws:s3:::bucket_name
    return bucket_arn.split(':')[-1]

# Function to get data source S3 configuration
def fetch_data_source_s3_configuration(data_source_id, knowledge_base_id):
    response = agent_client.get_data_source(dataSourceId=data_source_id, knowledgeBaseId=knowledge_base_id)
    s3_config = response['dataSource']['dataSourceConfiguration']['s3Configuration']
    bucket_arn = s3_config['bucketArn']
    bucket_name = extract_bucket_name(bucket_arn)
    s3_config['bucketName'] = bucket_name  # Add bucket name to the configuration dictionary
    return s3_config

# Function to reset session settings when the mode changes
def reset_session():
    st.session_state['session_enabled'] = False
    st.session_state['session_id'] = None
    st.session_state['first_input_processed'] = False
    st.session_state['user_input'] = ""

# Fetching regional agent and knowledge base lists
agent_summaries = fetch_agents()
knowledge_base_summaries = fetch_knowledge_bases()

agent_name_to_id = {agent['agentName']: agent['agentId'] for agent in agent_summaries}
agent_name_list = list(agent_name_to_id.keys())

# Create a mapping of knowledge base name to ID
kb_name_to_id = {kb['name']: kb['knowledgeBaseId'] for kb in knowledge_base_summaries}
kb_name_list = list(kb_name_to_id.keys())

# Initialize session state for session enabled and session ID
if 'session_enabled' not in st.session_state:
    st.session_state['session_enabled'] = False
if 'session_id' not in st.session_state:
    st.session_state['session_id'] = None
if 'first_input_processed' not in st.session_state:
    st.session_state['first_input_processed'] = False
if 'user_input' not in st.session_state:
    st.session_state['user_input'] = ""
if 'kb_id' not in st.session_state:
    st.session_state['kb_id'] = None
if 'data_source_id' not in st.session_state:
    st.session_state['data_source_id'] = None

# Streamlit App Layout
st.title('Bedrock Insurance Agent')
st.subheader('Powered by coffee and Amazon Bedrock')
st.info("**PURPOSE:** Allow users to select between Agents and Knowledge Bases for Amazon Bedrock for their task automation and intelligent search use cases. ")
idp_logo = "bedrock_logo.png"
st.sidebar.image(idp_logo, width=300, output_format='PNG')

# User choice: Agent or Knowledge Base
st.sidebar.subheader('1. Select Service Type')
use_agent = st.sidebar.radio("Mode of Operation", ["Agent", "Knowledge Base"], on_change=reset_session)

# Dropdowns for selecting agent name, agent alias ID, knowledge base ID, and model ID (conditionally displayed)
if use_agent == "Agent":
    st.sidebar.subheader('2. Choose Your Agent')
    agent_name = st.sidebar.selectbox("Agent Name", agent_name_list)
    agent_id = agent_name_to_id[agent_name]

    # Fetch agent aliases based on selected agent ID
    agent_alias_summaries = fetch_agent_aliases(agent_id) if agent_id else []
    agent_alias_id_list = [alias['agentAliasId'] for alias in agent_alias_summaries]
    agent_alias_id = st.sidebar.selectbox("Agent Alias ID", agent_alias_id_list) if agent_alias_id_list else st.sidebar.selectbox("Agent Alias ID", [])

else:
    st.sidebar.subheader('2. Choose Your Knowledge Base')
    kb_name_to_id = {kb['name']: kb['knowledgeBaseId'] for kb in knowledge_base_summaries}
    kb_name_list = list(kb_name_to_id.keys())
    kb_name = st.sidebar.selectbox("Knowledge Base Name", kb_name_list)
    kb_id = kb_name_to_id[kb_name]
    st.session_state['kb_id'] = kb_id

st.sidebar.subheader('3. Select Model ID')
model_id_list = ["anthropic.claude-3-sonnet", "anthropic.claude-3-haiku", "anthropic.claude-v2:1"]
model_id_map = {
    "anthropic.claude-3-sonnet": "anthropic.claude-3-sonnet-20240229-v1:0",
    "anthropic.claude-3-haiku": "anthropic.claude-3-haiku-20240307-v1:0",
    "anthropic.claude-v2:1": "anthropic.claude-v2:1"
}

model_selection = st.sidebar.selectbox("Model ID", model_id_list)
model_id = model_id_map[model_selection]
model_arn = f"arn:aws:bedrock:{region}::foundation-model/{model_id}"

# Enable Session checkbox, only enabled after the first input is processed
st.sidebar.subheader('4. Session Setting')
st.sidebar.radio("Enable Session", [False, True], key="session_enabled")

# Optional: Select filter attribute
st.sidebar.subheader('5. (Optional) Filter Setting')
filter_attributes = ["None", "external", "internal"]
filter_attribute = st.sidebar.selectbox("Filter Attribute", filter_attributes)

# Streamlit File Preview Helper Methods
def show_csv(uploaded_file):
    st.subheader("CSV Preview")
    df = pd.read_csv(uploaded_file)
    st.write(df)

def extract_text_from_docx(uploaded_file):
    try:
        document = Document(uploaded_file)
        text = ""
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from .doc(x) file: {e}")
        return None

def convert_docx_to_html(docx_content):
    try:
        document = Document(io.BytesIO(docx_content))
        text = "\n".join([paragraph.text for paragraph in document.paragraphs])
        html_content = f"<p>{text}</p>"
        return html_content
    except Exception as e:
        st.error(f"Error converting .docx to HTML: {e}")
        return None

def show_doc(uploaded_file):
    st.subheader("Document Preview")
    text = extract_text_from_docx(uploaded_file)
    if text:
        st.write(text)
    else:
        st.error("Uploaded file is not a valid Word document.")

def show_docx(uploaded_file):
    st.subheader("Document Preview")
    file_name = uploaded_file.name.lower()
    if 'docx' in file_name:
        docx_content = uploaded_file.getvalue()
        html_result = convert_docx_to_html(docx_content)
        if html_result:
            st.markdown(html_result, unsafe_allow_html=True)
        else:
            st.error("Failed to convert .docx to HTML")

def show_excel(uploaded_file):
    try:
        df = pd.read_excel(uploaded_file)
        st.subheader("Excel Preview")
        st.write(df)
    except Exception as e:
        st.error(f"Error reading Excel file: {e}")

def show_html(uploaded_file):
    st.subheader("HTML Preview")
    html_content = uploaded_file.getvalue().decode("utf-8")
    st.markdown(html_content, unsafe_allow_html=True)

def show_md(uploaded_file):
    st.subheader("Markdown Preview")
    md_content = uploaded_file.getvalue().decode("utf-8")
    st.markdown(md_content)

def show_pdf(uploaded_file):
    st.subheader("PDF Preview")
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64.b64encode(uploaded_file.read()).decode("utf-8")}" width="100%" height="500"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

def show_text(uploaded_file):
    text = uploaded_file.getvalue().decode("utf-8")
    st.subheader("Text Preview")
    st.write(text)

def process_uploaded_file(uploaded_file):
    file_name = uploaded_file.name.lower()
    file_extension = file_name.split(".")[-1]
    file_contents = None

    if file_extension == "csv":
        show_csv(uploaded_file)
        file_contents = uploaded_file.getvalue()

    elif file_extension == "doc":
        show_doc(uploaded_file)
        doc_content = extract_text_from_docx(uploaded_file)
        file_contents = doc_content.encode("utf-8") if doc_content else None

    elif file_extension == "docx":
        show_docx(uploaded_file)
        docx_content = extract_text_from_docx(uploaded_file)
        file_contents = docx_content.encode("utf-8") if docx_content else None

    elif file_extension in ["htm", "html"]:
        show_html(uploaded_file)
        html_content = uploaded_file.getvalue().decode("utf-8")
        file_contents = html_content.encode("utf-8") if html_content else None

    elif file_extension == "md":
        show_md(uploaded_file)
        md_content = uploaded_file.getvalue().decode("utf-8")
        file_contents = md_content.encode("utf-8") if md_content else None

    elif file_extension == "pdf":
        show_pdf(uploaded_file)
        file_contents = uploaded_file.getvalue()

    elif file_extension == "txt":
        show_text(uploaded_file)
        file_contents = uploaded_file.getvalue()

    elif file_extension in ["xls", "xlsx"]:
        show_excel(uploaded_file)
        file_contents = uploaded_file.getvalue()

    else:
        st.error("Preview not available for this file type.")

    return file_contents

# Agents and Knowledge Bases API Helper Functions
def bedrock_query_knowledge_base(query):
    print(f"Knowledge Base query: {query}")

    prompt_template = """\n\nHuman: You will be acting as a helpful customer service representative named Ava (short for Amazon Virtual Assistant) working for AnyCompany. Provide a summarized answer using only 1 or 2 sentences. 
    Here is the relevant information in numbered order from our knowledge base: $search_results$
    Current time: $current_time$
    User query: $query$\n\nAssistant: """

    payload = {
        "input": {
            "text": query
        },
        "retrieveAndGenerateConfiguration": {
            "type": "KNOWLEDGE_BASE",
            "knowledgeBaseConfiguration": {
                "generationConfiguration": {
                    "promptTemplate": {
                        "textPromptTemplate": prompt_template
                    }
                },
                "knowledgeBaseId": kb_id,
                "modelArn": model_arn,
                "retrievalConfiguration": {
                    "vectorSearchConfiguration": {
                        "numberOfResults": 5,
                    }
                }
            }
        }
    }

    if filter_attribute != "None":
        print(f"filter_attribute: {filter_attribute}")
        payload["retrieveAndGenerateConfiguration"]["knowledgeBaseConfiguration"]["retrievalConfiguration"] = {
            "vectorSearchConfiguration": {
                "numberOfResults": 5,
                "filter": {
                    "equals": {
                        "key": "exposure",
                        "value": filter_attribute
                    }
                }
            }
        }

    if st.session_state.get("session_enabled", True):
        if st.session_state["session_id"] is not None:
            sesh = st.session_state["session_id"]
            print(f"session_id: {sesh}")
            payload["sessionId"] = st.session_state["session_id"]

    try:
        response = agent_runtime_client.retrieve_and_generate(**payload)
        st.session_state["session_id"] = response.get("sessionId", st.session_state["session_id"])
        retrieval_results = response.get("retrievalResults", [])

        if 'output' in response:
            kb_response = response['output']['text']
            output_lines = kb_response.split('\n')
            clean_response = output_lines[-1].strip()
            print(f"Knowledge Base response ({model_selection}): {clean_response}\n")
            return clean_response
        else:
            return "No relevant information found in the knowledge base."

    except Exception as e:
        return f"Error querying knowledge base: {e}"

def update_knowledge_base(file_content, bucket_name, s3_file_name, selected_ds_id, selected_kb_id):
    print("Syncing Knowledge Base Data Source")

    try:
        file_obj = io.BytesIO(file_content)
        s3_client.upload_fileobj(file_obj, bucket_name, s3_file_name)
        st.success(f"File uploaded successfully to S3 bucket '{bucket_name}' as '{s3_file_name}'")
    except Exception as e:
        st.error(f"Error uploading file to S3: {e}")
        return

    description = "Programmatic update of Bedrock Knowledge Base Data Source"
    try:
        response = agent_client.start_ingestion_job(
            dataSourceId=selected_ds_id,
            description=description,
            knowledgeBaseId=selected_kb_id
        )
    except Exception as e:
        st.error(f"Error starting ingestion job: {e}")
    finally:
        file_obj.close()

def check_ingestion_job_status(selected_ds_id, selected_kb_id):
    headers = {
        "Content-type": "application/json",
    }

    status = ""
    while status != "complete":
        try:
            response = agent_client.list_ingestion_jobs(
                knowledgeBaseId=selected_kb_id,
                dataSourceId=selected_ds_id,
            )
            
            if response['ResponseMetadata']['HTTPStatusCode'] == 200:
                job_status = response["ingestionJobSummaries"][0]["status"]
                print(f"Ingestion Job Status: {job_status}")
                st.write(f"Ingestion Job Status: {job_status}")
                
                if job_status == "COMPLETE":
                    break
            else:
                st.write(f"Error: {response.status_code} - {response.text}")
        except Exception as e:
            st.write(f"An error occurred: {e}")

        time.sleep(4)

def invoke_agent(query):
    print(f"Agent query: {query}")

    # Generate new session ID if session is not enabled or session ID is None
    if not st.session_state["session_enabled"]:
        st.session_state["session_id"] = generate_session_id()
    
    try:    
        # Build the parameters for the invoke_agent call
        params = {
            'agentAliasId': agent_alias_id,
            'agentId': agent_id,
            'sessionId': st.session_state["session_id"],
            'inputText': query,
            'enableTrace': True
        }
        
        # Invoke the agent
        response = agent_runtime_client.invoke_agent(**params)
        
        # Handle the response (this is a simplification, adjust as needed for your use case)
        for event in response['completion']:
            if 'chunk' in event:
                result_bytes = event['chunk']['bytes']
                result_text = result_bytes.decode('utf-8')
                print(f"Agent response ({model_selection}): {result_text}\n")
                return result_text

    except Exception as e:
        return f"Error invoking agent: {e}"

def main():
    if not "valid_inputs_received" in st.session_state:
        st.session_state["valid_inputs_received"] = False

    if "uploaded_files" not in st.session_state:
        st.session_state["uploaded_files"] = []

    if use_agent == "Agent":
        st.subheader("Amazon Bedrock Agents - Prompt Input")
    else:
        st.subheader("Knowledge Bases for Amazon Bedrock - Prompt Input")
        selected_kb_id = kb_id
    
    query = st.text_input("User Input", value="", placeholder="What can the agent help you with?", label_visibility="visible")

    response = None
    if st.session_state.get("previous_query") != query and query != "":
        st.session_state['first_input_processed'] = True
        st.session_state["previous_query"] = query

        if use_agent == "Agent":
            response = invoke_agent(query)
        else:
            response = bedrock_query_knowledge_base(query)

    if response:
        st.write("Response:", response)

    st.subheader("Knowledge Bases for Amazon Bedrock - File Upload")

    if use_agent == "Agent":
        try:
            knowledge_bases = fetch_agent_knowledge_bases(agent_id)
            kb_ids = [kb['knowledgeBaseId'] for kb in knowledge_bases]
            kb_options = [fetch_knowledge_base_name(kb_id) for kb_id in kb_ids]

            selected_kb_name = st.selectbox("Select Knowledge Base", options=kb_options)
            selected_kb_id = kb_ids[kb_options.index(selected_kb_name)]

        except Exception as e:
            st.error(f"Error fetching knowledge bases: {str(e)}")

    # Fetch data sources and their IDs and names
    data_sources = fetch_data_sources(selected_kb_id)
    ds_options = [ds['name'] for ds in data_sources]
    selected_ds_name = st.selectbox("Select Data Source", options=ds_options)
    selected_ds_id = next(ds['id'] for ds in data_sources if ds['name'] == selected_ds_name)
    
    st.session_state['kb_id'] = selected_kb_id
    st.session_state['data_source_id'] = selected_ds_id
    s3_configuration = fetch_data_source_s3_configuration(selected_ds_id, selected_kb_id)

    uploaded_files = st.file_uploader("Upload Document", type=["csv", "doc", "docx", "htm", "html", "md", "pdf", "txt", "xls", "xlsx"], accept_multiple_files=True)

    if uploaded_files:
        for uploaded_file in uploaded_files:
            if uploaded_file not in st.session_state["uploaded_files"]:
                st.session_state["uploaded_files"].append(uploaded_file)
                file_name = "agent/knowledge-base-assets/" + uploaded_file.name
                file_contents = process_uploaded_file(uploaded_file)
                update_knowledge_base(file_contents, s3_configuration['bucketName'], file_name, selected_ds_id, selected_kb_id)
                check_ingestion_job_status(selected_ds_id, selected_kb_id)

if __name__ == "__main__":
    main()
